import os
import tempfile
import streamlit as st
from docx import Document
from docx.shared import Pt
from youtube_transcript_api import YouTubeTranscriptApi
import google.generativeai as genai

from pydub import AudioSegment

# -----------------------------
# UI Config
# -----------------------------
st.set_page_config(page_title="Audio/Video/YouTube → Transcript DOCX", layout="wide")
st.title("🎙️ Audio/Video/YouTube → Transcript + Notes → DOCX")
st.caption("Upload audio/video OR paste a YouTube link. Gemini API se structured transcript/notes banenge, phir DOCX download.")

# -----------------------------
# Helpers
# -----------------------------
def extract_youtube_id(url: str) -> str | None:
    # Simple extraction for common formats
    # Works for: https://www.youtube.com/watch?v=ID , https://youtu.be/ID
    import re
    patterns = [
        r"v=([a-zA-Z0-9_-]{6,})",
        r"youtu\.be/([a-zA-Z0-9_-]{6,})",
        r"youtube\.com/shorts/([a-zA-Z0-9_-]{6,})",
    ]
    for p in patterns:
        m = re.search(p, url)
        if m:
            return m.group(1)
    return None

def chunk_text(text: str, max_chars: int = 12000):
    # Gemini request size control
    chunks = []
    cur = []
    cur_len = 0
    for line in text.splitlines():
        if cur_len + len(line) + 1 > max_chars:
            chunks.append("\n".join(cur))
            cur = [line]
            cur_len = len(line) + 1
        else:
            cur.append(line)
            cur_len += len(line) + 1
    if cur:
        chunks.append("\n".join(cur))
    return chunks

def make_docx(title: str, content: str) -> bytes:
    doc = Document()

    # Title styling
    doc.add_heading(title, level=0)

    # Basic parsing: treat markdown-ish headings as doc headings
    for line in content.splitlines():
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph("")
            continue

        # Headings: "# ", "## ", "### "
        if line_stripped.startswith("### "):
            doc.add_heading(line_stripped.replace("### ", ""), level=3)
        elif line_stripped.startswith("## "):
            doc.add_heading(line_stripped.replace("## ", ""), level=2)
        elif line_stripped.startswith("# "):
            doc.add_heading(line_stripped.replace("# ", ""), level=1)
        else:
            # Bullets: "- " or "* "
            if line_stripped.startswith("- ") or line_stripped.startswith("* "):
                p = doc.add_paragraph(line_stripped[2:], style="List Bullet")
            else:
                p = doc.add_paragraph(line_stripped)

            # Slightly nicer font
            for run in p.runs:
                run.font.size = Pt(11)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    with open(tmp_path, "rb") as f:
        data = f.read()

    os.remove(tmp_path)
    return data

def gemini_generate_structured(api_key: str, raw_text: str, style: str, language: str):
    genai.configure(api_key=api_key)

    # Model: You can switch to latest available in your account
    model = genai.GenerativeModel("gemini-1.5-flash")

    # Prompt
    system_prompt = f"""
You are a transcription editor and technical note-writer.
You will receive raw transcript text.
Return a cleaned, structured output in Markdown.

Requirements:
- Language: {language}
- Output style: {style}
- Use clear headings and subheadings.
- Include:
  1) Title (as # ...)
  2) Key Takeaways (bullets)
  3) Detailed Transcript (with timestamps if present)
  4) Summary
- Fix filler words, improve readability, but don't change meaning.
- If transcript seems incomplete, mention it politely.
"""

    chunks = chunk_text(raw_text, max_chars=12000)

    outputs = []
    for i, ch in enumerate(chunks, start=1):
        prompt = system_prompt + f"\n\n[PART {i}/{len(chunks)}]\n\nRAW TRANSCRIPT:\n{ch}"
        resp = model.generate_content(prompt)
        outputs.append(resp.text)

    if len(outputs) == 1:
        return outputs[0]

    # Merge parts
    merge_prompt = f"""
Merge these parts into ONE coherent Markdown document.
Remove duplicates, keep best structure, ensure consistent headings.

Parts:
{chr(10).join([f"--- PART {idx+1} ---\n{t}" for idx, t in enumerate(outputs)])}
"""
    merged = model.generate_content(merge_prompt)
    return merged.text

def convert_to_wav(uploaded_file) -> str:
    # Save to temp then convert to wav (mono 16k if possible)
    suffix = os.path.splitext(uploaded_file.name)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        in_path = tmp.name

    # Convert using pydub (ffmpeg required)
    audio = AudioSegment.from_file(in_path)
    audio = audio.set_channels(1).set_frame_rate(16000)

    out_fd, out_path = tempfile.mkstemp(suffix=".wav")
    os.close(out_fd)
    audio.export(out_path, format="wav")

    os.remove(in_path)
    return out_path

def youtube_transcript_to_text(youtube_url: str) -> str:
    vid = extract_youtube_id(youtube_url)
    if not vid:
        raise ValueError("Invalid YouTube link. Video ID not found.")

    transcript = YouTubeTranscriptApi.get_transcript(vid)
    # transcript list: [{'text':..., 'start':..., 'duration':...}]
    lines = []
    for t in transcript:
        start = t.get("start", 0.0)
        txt = t.get("text", "").replace("\n", " ")
        lines.append(f"[{start:0.2f}s] {txt}")
    return "\n".join(lines)

# -----------------------------
# Sidebar Inputs
# -----------------------------
st.sidebar.header("⚙️ Settings")

api_key = st.sidebar.text_input("Gemini API Key", type="password", help="Apna Gemini API key yahan paste karo.")
language = st.sidebar.selectbox("Output Language", ["Hindi", "English", "Hinglish"], index=2)

style = st.sidebar.selectbox(
    "Document Style",
    ["Professional Notes", "Lecture Notes", "Meeting Minutes", "Podcast Show Notes"],
    index=0
)

st.sidebar.divider()
st.sidebar.caption("Tip: Community Cloud pe YouTube audio download restrictions ho sakti hain.")

# -----------------------------
# Main UI
# -----------------------------
tab1, tab2 = st.tabs(["📁 Upload Audio/Video", "▶️ YouTube Link"])

raw_text = None

with tab1:
    st.subheader("Upload Audio/Video")
    uploaded = st.file_uploader("Upload audio/video file", type=["mp3", "wav", "m4a", "mp4", "mov", "mkv"])
    st.info("Audio/video ko WAV (mono 16k) me convert karke text input flow me bheja jaayega (transcript aap provide kar sakte ho).")

    st.markdown("### Transcript Source")
    st.write("Aap ke paas transcript already ho to yahan paste kar do (best).")
    pasted = st.text_area("Paste raw transcript (optional)", height=200, placeholder="Yahan transcript paste karo...")

    if pasted.strip():
        raw_text = pasted.strip()
    elif uploaded:
        st.warning(
            "File upload ke saath automatic speech-to-text yahan include nahi kiya, "
            "kyunki Gemini STT approach different ho sakta hai. "
            "Aap chaho to main next step me Whisper / Google STT add kar dunga."
        )
        # Still convert to wav so user can extend later
        wav_path = convert_to_wav(uploaded)
        st.success("File converted to WAV. (Internal temp file ready)")
        st.write(f"Temp WAV path: `{wav_path}` (Streamlit runtime)")

with tab2:
    st.subheader("YouTube Link")
    yt = st.text_input("Paste YouTube URL", placeholder="https://www.youtube.com/watch?v=...")
    if yt:
        try:
            raw_text = youtube_transcript_to_text(yt)
            st.success("YouTube transcript fetched successfully.")
            with st.expander("Preview transcript"):
                st.text(raw_text[:4000] + ("..." if len(raw_text) > 4000 else ""))
        except Exception as e:
            st.error(f"Transcript fetch failed: {e}")
            st.caption("Agar video transcript available nahi hai, to aap manually transcript paste karke use kar sakte ho.")

# -----------------------------
# Generate
# -----------------------------
st.divider()
st.subheader("🧠 Generate Structured Transcript + DOCX")

col1, col2 = st.columns([1, 1])

with col1:
    doc_title = st.text_input("Document Title", value="Transcript & Notes")
with col2:
    file_name = st.text_input("Download File Name", value="transcript_notes.docx")

generate = st.button("Generate DOCX", type="primary", use_container_width=True)

if generate:
    if not api_key:
        st.error("Gemini API Key required.")
        st.stop()
    if not raw_text or len(raw_text.strip()) < 20:
        st.error("Please provide a transcript (paste) or fetch YouTube transcript.")
        st.stop()

    with st.spinner("Gemini se structured document ban raha hai..."):
        try:
            structured = gemini_generate_structured(api_key, raw_text, style=style, language=language)
        except Exception as e:
            st.error(f"Gemini generation failed: {e}")
            st.stop()

    st.success("Structured output ready!")

    st.markdown("### Preview (Markdown)")
    st.markdown(structured)

    docx_bytes = make_docx(doc_title, structured)

    st.download_button(
        label="⬇️ Download DOCX",
        data=docx_bytes,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
